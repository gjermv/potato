# -*- coding: UTF-8 -*-

'''
Created on 28. okt. 2019

@author: A485753
'''

from docx import Document
from docx.shared import Cm
import pandas as pd
from PIL import Image, ExifTags
import glob as glob
import os




def createPage(document, dict, image_folder, tunnelinfo={'tunnel_name': 'Eksempeltunnel'}):
    print(dict)
    document.add_heading('Bilag 3 - Registreringsskjema for skadested', 1)
    document.add_paragraph('Tunnelnavn: {}'.format(tunnelinfo['tunnel_name']))
    
    table1 = document.add_table(rows=3, cols=4)
    table1.style = document.styles['Light Grid']
    cell1 = table1.cell(0,3).merge(table1.cell(2,3))
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Retning:'
    hdr_cells[1].text = 'Prof.nr'
    hdr_cells[2].text = 'Plassering:'
    body_cells = table1.rows[1].cells
    body_cells[0].text = dict['Retning']
    body_cells[1].text = str(dict['Profnr'])
    

    paragraphX = cell1.paragraphs[0]
    run = paragraphX.add_run()
    
    
    if dict['Plassering'] == 'hs_vegg':
        body_cells[2].text = 'Høyre side, vegg'
        run.add_picture('C:\\python_proj\\Bilag3\\plassering\\hs_vegg.png', width=Cm(5))
        
    elif dict['Plassering'] == 'vs_vegg':
        body_cells[2].text = 'Venstre side, vegg'
        run.add_picture('C:\\python_proj\\Bilag3\\plassering\\vs_vegg.png', width=Cm(5))        
        
    elif dict['Plassering'] in ['vs_vederlag','vs_vl']:
        body_cells[2].text = 'Venstre side, vederlag'
        run.add_picture('C:\\python_proj\\Bilag3\\plassering\\vs_vl.png', width=Cm(5))        
        
    elif dict['Plassering'] in ['hs_vederlag','hs_vl']:
        body_cells[2].text = 'Høyre side, vederlag'
        run.add_picture('C:\\python_proj\\Bilag3\\plassering\\hs_vl.png', width=Cm(5))           
    
    elif dict['Plassering'] == 'heng':
        body_cells[2].text = 'Heng'
        run.add_picture('C:\\python_proj\\Bilag3\\plassering\\heng.png', width=Cm(5))                
    
    else:
        body_cells[1].text = 'Ukjent plassering'
        run.add_picture('C:\\python_proj\\Bilag3\\plassering\\ukjent2.png', width=Cm(5))             
    
       
    document.add_paragraph('', style='Normal')
    
    table2 = document.add_table(rows=2, cols=4)
    table2.style = document.styles['Light Grid']
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = 'Registering:'
    hdr_cells[1].text = 'Nedfall'
    hdr_cells[2].text = 'Plassforhold:'
    hdr_cells[3].text = 'Tiltak / Skadegrad'
    
    body_cells = table2.rows[1].cells
    body_cells[0].text = dict['Registering']
    body_cells[1].text = dict['Nedfall']
    body_cells[2].text = dict['Plassforhold']
    body_cells[3].text = str(dict['Tiltak'])
    
    document.add_paragraph('', style='Normal')
    table3 = document.add_table(rows=3, cols=1)
    table3.style = document.styles['Table Grid']
    hdr_cell = table3.rows[0].cells
    hdr_cell[0].text = 'Beskrivelse:'
    body_cell = table3.rows[1].cells
    body_cell[0].text = dict['Beskrivelse']
    
    image_cell = table3.rows[2].cells
    image_cell[0]
    
    paragraphY = image_cell[0].paragraphs[0]
    run2 = paragraphY.add_run()
    
    run2.add_picture('{}\\{}'.format(image_folder, dict['Bildenavn']), width=Cm(14.89))
    
    document.add_page_break()

def create_bilag3(tunnel_name, excel_fil, image_folder):
    document = Document()
    
    df = pd.read_excel(excel_fil, header=3)
    df_dict = df.to_dict(orient="index")
    print(df_dict)
    
    for key in df_dict:
        createPage(document, df_dict[key],image_folder,{'tunnel_name': tunnel_name})
    
    docName = 'C:\\python_proj\\Bilag3\\Bilag 3 - {}.docx'.format(tunnel_name)
    document.save(docName)
    
    return docName

def split_textTimestamp(myTxt):
    myData = myTxt.split(' ')
    return [myData[0],myData[1], myData[2]]



def TimeStampToExcel(folder):
    for image in glob.glob(folder+'\\*.jpg'):
        img = Image.open(image)
        exif_data = img._getexif()
        # print(exif_data)
        
        # for key, val in exif_data.items():
        #     if key in ExifTags.TAGS:
        #         print(key,f'{ExifTags.TAGS[key]}:{val}')
        txt = exif_data[270].replace('ASCII','').strip('\x00').strip()
        
        txt = txt.replace('\n', ' ').replace('Ã¸','ø').replace('Ã¥','å')
        #txt_tab = txt.split(' ',2)
        print(txt)
        #print(txt_tab[0].lower(),txt_tab[1],txt_tab[2],os.path.basename(image).split('.')[0],sep=';')

def insertPDFToDWG():
    folder = 'C:\\Users\\A485753\\OneDrive - AFRY\\E18 Arbeidskatalog\\10 Fag\\RIG\\Ing.geo\\Parsell 3\\2022-03-10 Tunnelinspeksjon Bamble\\02 Kartleggingsskjema\\'
    file = 'Kartleggingsskjema Bamble scannet.pdf'
    filename = folder + file
    profile_start = 16200
    
    for i in range(6):
        print("""-ATTACH {}
{}
0,0
1
0
AL
F
10,5
-1,5


1.0534,1.0079
-44,{}
5.2034,10.3147
44,{}

y""".format(filename, i+1,i*200+profile_start,i*200+200+profile_start)    )


if __name__ == '__main__':
    TimeStampToExcel('C:\\Users\\A485753\\OneDrive - AFRY\\E18 Arbeidskatalog\\10 Fag\\RIG\\Ing.geo\\Parsell 3\\2022-03-10 Tunnelinspeksjon Bamble\\01 Bilde')
    #insertPDFToDWG()
    # print(create_bilag3('Espatunnelen - nordgående', 'C:\python_proj\\Bilag3\\Test_data.xlsx','C:\\python_proj\\Bilag3\\Test_data'))
    
#     df = pd.read_excel('T:\\19500 - Tunnelinspeksjoner SVV Region Øst 2019\\018 Follotunnelen\\04 Rapport\\01 Arbeidsmappe\\Bilag 3\\Bilag3_FolloSør.xlsx', header=3)
#     df_dict = df.to_dict(orient="index")
#     print(df_dict)
#     
#     for key in df_dict:
#         createPage(document, df_dict[key],{'tunnel_name': 'Follotunnelen - sørgående løp'})
#     #document.save('C:\\python_proj\\Bilag3\\Bilag 3 - Skaderegistrering Oslofjordtunnelen.docx')
#     docName = 'C:\\python_proj\\Bilag3\\Bilag 3 - Follotunnelen - Sørgående løp.docx'
#     document.save(docName)
#     
#     print('Document saved: {}'.format(docName))